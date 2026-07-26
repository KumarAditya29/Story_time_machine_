from __future__ import annotations

from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "StoryTime_Machine_Technical_Documentation.docx"
ASSETS = ROOT / "docs" / "_doc_assets"
ASSETS.mkdir(parents=True, exist_ok=True)

INK = "14213D"
PURPLE = "7653D6"
LAVENDER = "EDE9FE"
MUTED = "5E6B85"
PALE = "F4F6FB"
GREEN = "147D64"
RED = "A9374B"


def font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def box(draw, xy, title, subtitle="", fill="#1B2440", outline="#7653D6"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    draw.text((x1 + 18, y1 + 16), title, font=font(24, True), fill="white")
    if subtitle:
        draw.multiline_text((x1 + 18, y1 + 51), subtitle, font=font(16), fill="#CAD3EB", spacing=4)


def arrow(draw, start, end, label=""):
    draw.line([start, end], fill="#9AA7C4", width=4)
    x, y = end
    draw.polygon([(x, y), (x - 12, y - 7), (x - 12, y + 7)], fill="#9AA7C4")
    if label:
        mx, my = (start[0] + end[0]) // 2, (start[1] + end[1]) // 2
        draw.rounded_rectangle((mx - 46, my - 16, mx + 46, my + 14), radius=7, fill="#EEF1F8")
        draw.text((mx - 35, my - 10), label, font=font(13, True), fill="#44516E")


def system_diagram(path: Path):
    img = Image.new("RGB", (1500, 800), "#0C1020")
    d = ImageDraw.Draw(img)
    d.text((55, 35), "StoryTime Machine — High-Level System Architecture", font=font(34, True), fill="white")
    box(d, (70, 220, 350, 390), "Writer", "Selects story • edits scene\nreviews impact and diffs", "#25304E")
    box(d, (500, 160, 860, 450), "React Web App", "Story library and accordion\nGraph canvas (React Flow)\nScene editor • change review", "#2D2252")
    box(d, (1010, 120, 1400, 280), "FastAPI API", "Thin HTTP boundary\nCORS • request validation\nStory / graph / edit endpoints", "#25304E")
    box(d, (1010, 360, 1200, 560), "Core Engine", "NetworkX graph\nVersioning & audit\nJSON storage", "#1D4A50", "#42B9A2")
    box(d, (1230, 360, 1420, 560), "OpenAI", "Structured state extraction\nVerification\nMinimal rewrite", "#4A2A50", "#D983D7")
    arrow(d, (350, 305), (500, 305), "HTTP")
    arrow(d, (860, 240), (1010, 200), "REST")
    arrow(d, (1205, 280), (1110, 360), "calls")
    arrow(d, (1210, 450), (1230, 450), "LLM")
    img.save(path)


def data_diagram(path: Path):
    img = Image.new("RGB", (1500, 850), "#0C1020")
    d = ImageDraw.Draw(img)
    d.text((55, 35), "Story Model and Dependency Graph", font=font(34, True), fill="white")
    box(d, (70, 160, 400, 390), "Story JSON", "story_id • title • genre\nscenes[] • dependencies[]\nversions[] • timestamps", "#25304E")
    box(d, (540, 115, 925, 350), "Atomic Scene", "scene_id • title • text • order\nStoryBible metadata\ncharacters • facts • reads/writes\ncausal setup/payoff • tone", "#2D2252")
    box(d, (1040, 115, 1410, 350), "Dependency", "source_scene_id → target_scene_id\nkind: fact / character / causal\nreason • confidence", "#1D4A50", "#42B9A2")
    box(d, (540, 470, 925, 680), "StoryVersion", "parent_version_id • label\nchanged_scene_ids\naudit events • snapshots", "#4A2A50", "#D983D7")
    box(d, (1040, 470, 1410, 680), "AuditEntry", "edited / checked / regenerated / skipped\nhop • reason • confidence\nbefore_text • after_text", "#3B344E")
    arrow(d, (400, 265), (540, 235), "contains")
    arrow(d, (925, 235), (1040, 235), "links")
    arrow(d, (270, 390), (620, 470), "versions")
    arrow(d, (925, 575), (1040, 575), "records")
    img.save(path)


def flow_diagram(path: Path):
    img = Image.new("RGB", (1500, 960), "#0C1020")
    d = ImageDraw.Draw(img)
    d.text((55, 35), "Edit → Verify → Regenerate Sequence", font=font(34, True), fill="white")
    steps = [
        ("1. Writer edits scene", "New text is submitted; a new version begins."),
        ("2. Extract new story state", "OpenAI creates the updated scene bible."),
        ("3. Diff old vs. new state", "Facts, character state, tone, and causal changes are identified."),
        ("4. Select candidates", "BFS over union of pre-edit + post-edit graph; visited set prevents cycles."),
        ("5. Verify each candidate", "LLM answers whether consistency actually requires a change."),
        ("6. Minimal rewrite", "Only confirmed scenes are rewritten with neighbor voice context."),
        ("7. Re-extract & persist", "Graph, audit trail, snapshots, and UI diffs are updated."),
    ]
    y = 120
    for index, (title, subtitle) in enumerate(steps):
        box(d, (245, y, 1255, y + 90), title, subtitle, "#25304E" if index not in (4, 5) else "#4A2A50")
        if index < len(steps) - 1:
            arrow(d, (750, y + 90), (750, y + 122))
        y += 122
    img.save(path)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            set_cell_margins(cell)


def style_cell(cell, header=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
            run.font.size = Pt(9.5 if not header else 10)
            run.font.color.rgb = RGBColor.from_string("FFFFFF" if header else INK)
            run.bold = header
    if header:
        shade(cell, PURPLE)
    else:
        shade(cell, "FFFFFF")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        style_cell(header_cells[idx], True)
    set_repeat_table_header(table.rows[0])
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = value
            style_cell(cells[idx])
        prevent_row_split(table.rows[-1])
    set_table_widths(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullets(doc, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)


def add_code(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    shade_paragraph(paragraph, "F0F2F8")
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string("24314C")


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def build():
    system = ASSETS / "system_architecture.png"
    data = ASSETS / "data_model.png"
    flow = ASSETS / "edit_flow.png"
    system_diagram(system)
    data_diagram(data)
    flow_diagram(flow)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color, before, after in [("Heading 1", 16, PURPLE, 15, 7), ("Heading 2", 13, "2E4A7D", 11, 5), ("Heading 3", 11.5, INK, 8, 3)]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    # Header/footer
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("STORYTIME MACHINE  |  TECHNICAL DOCUMENTATION")
    hr.font.name = "Calibri"; hr.font.size = Pt(8); hr.font.bold = True; hr.font.color.rgb = RGBColor.from_string(MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("Pocket FM Hackathon Project  •  Local Demo Build  •  26 July 2026")
    fr.font.name = "Calibri"; fr.font.size = Pt(8); fr.font.color.rgb = RGBColor.from_string(MUTED)

    # Cover
    doc.add_paragraph().paragraph_format.space_after = Pt(35)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(10)
    run = kicker.add_run("POCKET FM HACKATHON • TECHNICAL DESIGN DOCUMENT")
    run.font.name = "Calibri"; run.font.size = Pt(10); run.bold = True; run.font.color.rgb = RGBColor.from_string(PURPLE)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("StoryTime Machine")
    run.font.name = "Calibri"; run.font.size = Pt(34); run.bold = True; run.font.color.rgb = RGBColor.from_string(INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(24)
    run = subtitle.add_run("Narrative version control with dependency-aware AI regeneration")
    run.font.name = "Calibri"; run.font.size = Pt(16); run.font.color.rgb = RGBColor.from_string(MUTED)
    add_body(doc, "A local full-stack writer platform that models a story as a dependency graph, computes the causal blast radius of an edit, verifies every candidate scene, and rewrites only the scenes that truly need continuity repair.")
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    add_table(doc, ["Document", "Value"], [
        ("Version", "1.0 — implementation-aligned technical documentation"),
        ("Audience", "Hackathon judges, Pocket FM product/engineering reviewers, and developers"),
        ("Stack", "React + TypeScript, FastAPI + Python, NetworkX, OpenAI structured outputs, JSON storage"),
        ("Primary guarantee", "Untouched scenes remain byte-identical; all changes have an auditable reason and diff."),
    ], [1.4, 5.1])
    page_break(doc)

    # TOC-style navigation
    add_heading(doc, "Document Guide", 1)
    add_body(doc, "This document is designed as both an architecture reference and a hackathon demo guide. It describes the problem, implementation, data model, APIs, AI prompts, operational flow, user experience, and known constraints.")
    add_table(doc, ["Section", "Purpose"], [
        ("1. Product and problem", "Narrative drift problem, solution promise, scope, and success criteria."),
        ("2. Architecture", "System context, modules, component responsibilities, and deployment topology."),
        ("3. Data and graph engine", "Schemas, graph rules, blast-radius traversal, and version model."),
        ("4. AI consistency pipeline", "Extraction, metadata delta, verification, minimal rewrite, and multi-hop behavior."),
        ("5. Interface and API", "Writer experience, endpoints, request/response patterns, and state handling."),
        ("6. Operations", "Setup, test evidence, security, limitations, and demo runbook."),
    ], [1.6, 4.9])

    add_heading(doc, "1. Product and Problem", 1)
    add_heading(doc, "1.1 Narrative drift in conventional regeneration", 2)
    add_body(doc, "When a writer changes an early scene, conventional AI writing tools commonly regenerate every later scene. That approach is slow, costly, and risky: unrelated scenes can be reworded, voice can drift, and the writer cannot prove why any particular scene changed. Pocket FM-style serial fiction compounds this because character state, secrets, objects, and emotional promises persist across long story arcs.")
    add_heading(doc, "1.2 Solution thesis", 2)
    add_body(doc, "StoryTime Machine treats a story as an incremental build graph rather than a flat document. Atomic scenes publish narrative state and consume state established upstream. An edit creates a new version; the system uses dependency edges to select a minimal blast radius, verifies each candidate, and rewrites only confirmed inconsistencies. The resulting audit provides an evidence trail for every decision.")
    add_heading(doc, "1.3 Functional scope", 2)
    add_bullets(doc, [
        "Ten pre-seeded Pocket FM-style stories across multiple genres, with ten scenes each (100 scenes total).",
        "Create a new story from a title, genre, logline, and multiple pasted scenes.",
        "Map a story into structured scene metadata using the OpenAI API.",
        "Visualize explicit scene-to-scene narrative dependencies.",
        "Edit any story, any scene, any number of times; every save creates a new version.",
        "Review before/after differences for writer edits and AI minimal rewrites.",
        "Inspect a scene accordion from the story library and navigate directly to any scene.",
    ])
    add_heading(doc, "1.4 Non-goals for this hackathon build", 2)
    add_bullets(doc, [
        "Multi-user collaboration, authentication, and cloud-hosted persistence.",
        "Full natural-language retrieval across an entire catalog.",
        "Automatic publication to Pocket FM or a production editorial CMS.",
        "Immutable Git-style branch switching in the UI (versions are persisted and auditable; a restore UI is a planned extension).",
    ])

    add_heading(doc, "2. High-Level System Design", 1)
    doc.add_picture(str(system), width=Inches(6.65))
    cap = doc.add_paragraph("Figure 1. High-level context: the React writer workspace calls a thin FastAPI API, which coordinates the graph/versioning modules, JSON persistence, and OpenAI.")
    cap.paragraph_format.space_after = Pt(8); cap.runs[0].italic = True; cap.runs[0].font.size = Pt(9); cap.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    add_heading(doc, "2.1 Request path", 2)
    add_bullets(doc, [
        "The writer interacts only with the React frontend at http://127.0.0.1:5173.",
        "The frontend calls the FastAPI backend at http://127.0.0.1:8000/api.",
        "FastAPI validates input, calls local Python modules directly (no internal REST hops), and returns JSON.",
        "OpenAI is used only by the backend; the API key is loaded from a local ignored .env file and is never exposed to the client.",
        "Each story is serialized as an independent JSON file, making the demo portable and inspectable.",
    ])
    add_heading(doc, "2.2 Technology choices", 2)
    add_table(doc, ["Layer", "Technology", "Reason"], [
        ("Frontend", "React 18, TypeScript, Vite", "Fast interactive writer workspace and maintainable UI state."),
        ("Graph UI", "React Flow (@xyflow/react)", "Node-edge visualization, selection, controls, and dependency labels."),
        ("Backend", "Python 3, FastAPI, Pydantic", "Thin typed API wrapper around local orchestration modules."),
        ("Graph engine", "NetworkX", "Directed graph construction and bounded BFS traversal."),
        ("LLM", "OpenAI gpt-4o-mini", "Low-cost structured extraction, verification, and minimal rewrites."),
        ("Persistence", "One JSON file per story", "Transparent, portable, low-friction hackathon storage."),
    ], [1.0, 1.7, 3.8])

    add_heading(doc, "3. Low-Level Architecture", 1)
    add_heading(doc, "3.1 Repository structure", 2)
    add_code(doc, "backend/main.py            FastAPI routes and CORS\nbackend/models.py          Pydantic domain contracts\nbackend/storage.py         JSON repository (one file per story)\nbackend/seed_data.py       10-story / 100-scene demo corpus\nbackend/generation.py      OpenAI extraction, verification, rewrite\nbackend/graph_engine.py    NetworkX dependency construction and BFS\nbackend/versioning.py      Versioned edit orchestration and audit\nfrontend/src/main.tsx      React writer workspace\nfrontend/src/styles.css    Product UI, graph, accordion, diff styling\nGOALS.md                   Milestone roadmap\nREADME.md                  Setup and demo instructions")
    add_heading(doc, "3.2 Module responsibilities", 2)
    add_table(doc, ["Module", "Responsibility", "Key behavior"], [
        ("models.py", "Domain model", "Defines Scene, StoryBible, Dependency, AuditEntry, StoryVersion, and Story."),
        ("storage.py", "Persistence boundary", "Validates story IDs, reads/writes JSON, lists story summaries."),
        ("generation.py", "LLM adapter", "Uses structured outputs; bounds requests to 35 seconds with one retry."),
        ("graph_engine.py", "Causal graph", "Builds explainable edges and selects impact candidates with BFS."),
        ("versioning.py", "Edit transaction", "Creates snapshot/version, verifies candidates, regenerates confirmed scenes, persists audit."),
        ("main.py", "HTTP surface", "Routes, CORS, response typing, startup seed initialization, and error conversion."),
    ], [1.15, 1.45, 3.9])

    add_heading(doc, "4. Data Model and Persistence", 1)
    doc.add_picture(str(data), width=Inches(6.65))
    cap = doc.add_paragraph("Figure 2. Durable story model. A Story holds ordered scenes, graph edges, and append-only version/audit records.")
    cap.runs[0].italic = True; cap.runs[0].font.size = Pt(9); cap.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    add_heading(doc, "4.1 Atomic scene", 2)
    add_code(doc, '{\n  "scene_id": "s07",\n  "title": "The Concert Trap",\n  "text": "...",\n  "order": 7,\n  "bible": {\n    "characters_present": ["Tara", "Iqbal"],\n    "character_states": {"Tara": "determined"},\n    "objects_facts": ["the sitar contains footage"],\n    "causal_setup": ["concert trap"],\n    "reads": ["sitar_contains_footage"],\n    "writes": [{"state_var_id": "concert_trap_ready", "new_value": "true"}]\n  }\n}')
    add_heading(doc, "4.2 Story bible fields", 2)
    add_table(doc, ["Field", "Meaning", "Used by"], [
        ("characters_present", "Named participants in this scene.", "Character continuity candidates."),
        ("character_states", "Relevant relationship, belief, status, or trait.", "Continuity and metadata delta."),
        ("objects_facts / established_facts", "Props, secrets, and facts present or established.", "Fact dependency candidates."),
        ("causal_setup / causal_payoff_of", "Future promise and resolved earlier thread.", "High-confidence causal edges."),
        ("reads / writes", "Stable state IDs consumed or emitted by a scene.", "Fine-grained state graph and delta evidence."),
        ("emotional_tone / location", "Voice and spatial context.", "Rewrite consistency and delta explanation."),
    ], [1.55, 2.3, 2.65])
    add_heading(doc, "4.3 Version persistence", 2)
    add_body(doc, "A save does not overwrite history. A StoryVersion has a unique version_id, a parent_version_id, a human-readable label, changed scene IDs, an audit list, and snapshots of scenes/dependencies. This is the foundation for future branch comparison or restore features.")

    add_heading(doc, "5. Dependency Graph Engine", 1)
    add_heading(doc, "5.1 Edge construction", 2)
    add_table(doc, ["Edge type", "Rule", "Default confidence"], [
        ("fact", "A later scene reads/references a state or fact written/established by an earlier scene.", "0.90"),
        ("causal", "Earlier causal_setup matches later causal_payoff_of.", "0.96"),
        ("character", "Scenes share a character; stronger if both contain state for that character.", "0.50–0.72"),
        ("foreshadow", "Scenes share a recurring planted narrative thread.", "0.62"),
    ], [1.1, 4.2, 1.15])
    add_heading(doc, "5.2 Blast radius algorithm", 2)
    add_bullets(doc, [
        "Build a directed graph where nodes are scene IDs and edges carry kind, reason, and confidence.",
        "On edit, form the union of pre-edit and post-edit dependency edges. This is critical: a change that removes a fact must not disconnect the old consumers before they are checked.",
        "Run bounded breadth-first traversal with a visited set from the edited scene. Forward links find later effects; backward causal/foreshadow links preserve setup-payoff consistency.",
        "Limit propagation to three hops. Confidence thresholds become stricter for later hops (0.35, 0.62, 0.82) to prevent runaway cascades.",
        "Every candidate is passed to LLM verification before any rewrite is permitted.",
    ])
    add_heading(doc, "5.3 Why the graph is not simply chronological", 2)
    add_body(doc, "A story can contain independent subplots, scenes with the same character but no state dependency, or late scenes that do not rely on an early changed fact. Chronological regeneration is therefore over-inclusive. The graph has an explicit reason on each edge so the interface can show why a scene was considered and the audit can show why it was changed or skipped.")

    add_heading(doc, "6. AI Consistency and Regeneration Pipeline", 1)
    doc.add_picture(str(flow), width=Inches(6.65))
    cap = doc.add_paragraph("Figure 3. The edit transaction. Candidate selection is deterministic; AI is used for structure extraction, consistency judgment, and minimal text repair.")
    cap.runs[0].italic = True; cap.runs[0].font.size = Pt(9); cap.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    add_heading(doc, "6.1 Structured state extraction", 2)
    add_body(doc, "The extract_state function asks OpenAI for JSON conforming to a Pydantic schema. The prompt explicitly asks for reusable snake_case state identifiers, scene-local facts only, and causal setup/payoff. Character state is represented as a closed list at the API boundary to comply with strict JSON schema requirements, then converted to the durable dictionary model.")
    add_heading(doc, "6.2 Metadata delta", 2)
    add_body(doc, "After an edit, the engine compares the old and new StoryBible fields. The delta records additions/removals across character state, facts, reads/writes, emotional tone, and causal signals. This is the compact evidence passed to verification and rewrite prompts rather than flooding the model with the entire book.")
    add_heading(doc, "6.3 Candidate verification", 2)
    add_code(doc, "Given the original upstream scene, edited upstream scene, metadata delta, candidate scene, graph reason, and hop: does this candidate need to change to remain consistent? Return needs_change, reason, and the minimum required change.")
    add_body(doc, "The verifier is intentionally conservative. A graph relationship selects a scene for review, not automatic modification. A checked scene may be explicitly skipped with a reason, preserving byte-identical text and proving that it was considered.")
    add_heading(doc, "6.4 Minimal regeneration", 2)
    add_body(doc, "For a confirmed candidate, the rewrite prompt receives only the original candidate text, the relevant delta plus the verifier’s required change, and one neighboring scene on each side for voice context. The instruction prioritizes preserving wording, pacing, facts, and unaffected content. The rewritten scene is immediately re-extracted so later edits use current state.")
    add_heading(doc, "6.5 Timeouts and resiliency", 2)
    add_body(doc, "OpenAI client calls use a 35-second timeout and one retry. The frontend presents processing status for mapping and propagation. If a model request fails, FastAPI converts the failure to a 503 response and the UI displays an error instead of an indefinite spinner.")

    add_heading(doc, "7. API Design", 1)
    add_table(doc, ["Method", "Route", "Purpose"], [
        ("GET", "/api/health", "Backend health status."),
        ("GET", "/api/stories", "Story-library summaries."),
        ("POST", "/api/stories", "Create a story with multiple scenes."),
        ("GET", "/api/stories/{story_id}", "Load full story, scenes, graph, and versions."),
        ("POST", "/api/stories/{story_id}/analyze", "OpenAI story-bible extraction and graph construction."),
        ("GET", "/api/stories/{story_id}/graph", "Graph-ready node/edge representation."),
        ("POST", "/api/stories/{story_id}/scenes/{scene_id}/edit", "Create a versioned edit and propagate consistency changes."),
    ], [0.65, 2.75, 3.05])
    add_heading(doc, "7.1 Edit request", 2)
    add_code(doc, 'POST /api/stories/velvet-voicemail/scenes/s06/edit\n{\n  "text": "Mira finds the rings engraved with Vikram’s childhood nickname...",\n  "label": "Writer edit — The Engraving"\n}')
    add_heading(doc, "7.2 Edit response", 2)
    add_body(doc, "The response returns the complete updated Story. Its latest version contains changed_scene_ids and a chronological audit. UI state is updated from this response, allowing the graph, scene list, Change Review panel, and full before/after inspector to remain synchronized.")
    add_heading(doc, "7.3 API documentation", 2)
    add_body(doc, "During local execution, interactive Swagger documentation is available at http://127.0.0.1:8000/docs. The API accepts lowercase scene IDs such as s01, s02, and s03.")

    add_heading(doc, "8. Writer Interface Design", 1)
    add_heading(doc, "8.1 Major UI regions", 2)
    add_table(doc, ["Region", "Interaction", "Evidence surfaced"], [
        ("Story library", "Expand chevron for a scene accordion; select a story or direct scene.", "Story title, genre, scene count, and child scene titles."),
        ("Dependency map", "Click nodes; edges show fact/character/causal labels.", "Mapped status, writer edit, AI rewrite, dependency reasons."),
        ("Scene editor", "Edit any current scene, including one previously edited or regenerated.", "Editable text, character/location context, Save & propagate."),
        ("Scene diff inspector", "Open a changed scene and click Edit this scene again when needed.", "Visible old/new content and whether the change came from writer or AI."),
        ("Change Review", "Click a compact diff to focus its scene.", "All edit/check/regenerate/skip decisions, reason, hop, and old/new text."),
    ], [1.35, 2.45, 2.65])
    add_heading(doc, "8.2 Unlimited edit behavior", 2)
    add_body(doc, "The UI no longer treats a change as terminal. Every scene can be selected from the graph, main scene list, or story accordion. A previously changed scene displays its diff and an Edit this scene again action. The subsequent save creates another StoryVersion using the latest story state and runs the same graph/verification/rewrite transaction. Earlier versions are retained in JSON with parent relationships and snapshots.")
    add_heading(doc, "8.3 Visual language", 2)
    add_bullets(doc, [
        "Purple graph nodes and spark markers indicate a writer edit or AI rewrite in the current version.",
        "Dark purple dependency-label chips distinguish fact, character, causal, and foreshadow relationships without white default labels.",
        "Red OLD and green NEW panels make textual differences immediately visible.",
        "Checked/skipped entries remain visible to prove the system avoided unnecessary rewriting.",
    ])

    add_heading(doc, "9. Setup and Local Operation", 1)
    add_heading(doc, "9.1 Prerequisites", 2)
    add_bullets(doc, ["Python 3.11+", "Node.js 18+", "An OpenAI API key with access to the configured model."])
    add_heading(doc, "9.2 Environment", 2)
    add_code(doc, "cp .env.example .env\n# Set OPENAI_API_KEY in .env\n# OPENAI_MODEL defaults to gpt-4o-mini")
    add_heading(doc, "9.3 Start services", 2)
    add_code(doc, "python3 -m pip install -r requirements.txt\npython3 -m uvicorn backend.main:app --reload --port 8000\n\ncd frontend\nnpm install\nnpm run dev")
    add_body(doc, "Open http://127.0.0.1:5173 for the app. Do not open frontend/index.html with a file:// URL: Vite must serve the React application. Open http://127.0.0.1:8000/docs for Swagger endpoints.")
    add_heading(doc, "9.4 Seed library", 2)
    add_body(doc, "On backend startup, seed_data.seed_stories creates missing JSON stories. The library contains Velvet Voicemail (romance), Midnight Platform (supernatural thriller), Throne of Embers (fantasy), The Last Monsoon (climate mystery), Cipher in the Sitar (musical crime), Neon Inheritance (techno noir), Tea Estate Secret (family drama), Orbit of Us (space romance), Courtroom of Shadows (legal thriller), and The Echo Village (folklore horror).")

    add_heading(doc, "10. Testing and Evidence", 1)
    add_table(doc, ["Check", "Evidence"], [
        ("Seed validation", "Ten stories with at least ten ordered scenes each; 100 seed scenes total."),
        ("Graph unit smoke test", "An unrelated scene is excluded from a causal blast radius."),
        ("Live LLM extraction", "Structured extraction was exercised against seed scenes using the configured OpenAI model."),
        ("Live story analysis", "A complete ten-scene story was mapped and persisted with dependency edges."),
        ("Live edit flow", "A ring-inscription change checked downstream candidates and regenerated only confirmed scenes."),
        ("Frontend build", "TypeScript and Vite production build succeeded after UI enhancements."),
        ("Browser smoke test", "Story library, graph canvas, mapping action, and CORS behavior were verified locally."),
    ], [1.55, 4.9])
    add_code(doc, "python3 -m pytest -q\ncd frontend && npm run build")

    add_heading(doc, "11. Security, Privacy, and Cost", 1)
    add_bullets(doc, [
        "OPENAI_API_KEY is loaded only on the backend from .env; .env is ignored by Git and never returned to the browser.",
        "Story data is local JSON by default. No database, identity provider, or cloud storage is required for the demo.",
        "Only relevant scene text, adjacent voice context, and compact state delta are sent to OpenAI per operation; the entire story is not sent during a rewrite.",
        "gpt-4o-mini is selected to make per-scene extraction and candidate verification cost-conscious for a hackathon demo.",
        "Because a key was manually supplied for local testing, it should be rotated after the hackathon and replaced through a secure secret manager for production.",
    ])

    add_heading(doc, "12. Limitations and Production Roadmap", 1)
    add_table(doc, ["Current limitation", "Production-grade extension"], [
        ("JSON file persistence", "Use Postgres/SQLite with transactional version snapshots and object storage for long scene text."),
        ("Single-process analysis", "Use background jobs, progress events/WebSockets, cancellation, and rate-limit handling."),
        ("Heuristic edge construction", "Add an explicit LLM relationship pass, embeddings, human edge review, and graph quality scoring."),
        ("No UI restore/branch switch", "Expose version timeline, compare, restore, named branches, and merge/approval workflows."),
        ("No authentication", "Add writer roles, story ownership, editorial review permissions, and audit identity."),
        ("No production observability", "Add tracing, prompt/cost metrics, retry telemetry, and hallucination/consistency evaluation datasets."),
    ], [2.25, 4.2])

    add_heading(doc, "13. Hackathon Demo Runbook", 1)
    add_bullets(doc, [
        "Open the app and expand a story’s chevron to show its scene list; select an early pivot scene.",
        "Click Map story with AI. Explain that each scene receives a story-bible record and the graph becomes explainable rather than chronological.",
        "Choose Velvet Voicemail, Scene 6: The Engraving. Change the ring inscription from Mira’s mother’s initials to Vikram’s childhood nickname.",
        "Save & propagate. Explain the exact pipeline: metadata delta → pre/post graph union → bounded BFS → verification → minimal rewrite.",
        "Open Change Review. Show the writer’s old/new diff, every checked scene, regenerated scene diffs, reason, and hop number.",
        "Select a rewritten scene, show its full diff, then click Edit this scene again to demonstrate iterative versioning.",
        "Close with the guarantee: scenes not marked with a purple spark were not regenerated and remain byte-identical.",
    ])

    add_heading(doc, "14. Conclusion", 1)
    add_body(doc, "StoryTime Machine demonstrates that AI story editing does not need to be an opaque rewrite of everything downstream. By expressing narrative continuity as explicit state and dependencies, it provides a practical incremental-regeneration workflow: faster changes, less drift, explainable causality, and a versioned audit trail a writer can trust. The hackathon implementation is intentionally local and lightweight, but its contracts and modules are structured for direct evolution into an editorial platform for long-form serial fiction.")

    doc.core_properties.title = "StoryTime Machine — Technical Documentation"
    doc.core_properties.subject = "Architecture, design, and operation of narrative version control platform"
    doc.core_properties.author = "StoryTime Machine Team"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
